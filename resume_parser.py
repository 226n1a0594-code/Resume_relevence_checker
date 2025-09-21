import re
import spacy
from typing import Dict, List, Optional
import fitz  # PyMuPDF
import docx
from io import BytesIO
import streamlit as st

class ResumeParser:
    def __init__(self):
        self.skills_keywords = self._load_comprehensive_skills_database()
        try:
            self.nlp = spacy.load("en_core_web_sm")
        except OSError:
            st.warning("⚠️ spaCy model not found. Using basic parsing mode.")
            self.nlp = None
    
    def _load_comprehensive_skills_database(self) -> Dict[str, List[str]]:
        """Load comprehensive skills database with latest technologies"""
        return {
            'programming_languages': [
                'python', 'java', 'javascript', 'typescript', 'c++', 'c#', 'c', 'php', 'ruby', 
                'go', 'rust', 'swift', 'kotlin', 'scala', 'r', 'matlab', 'perl', 'lua', 'dart',
                'objective-c', 'assembly', 'cobol', 'fortran', 'haskell', 'erlang', 'elixir',
                'clojure', 'f#', 'visual basic', 'shell scripting', 'bash', 'powershell', 'zsh'
            ],
            'web_technologies': [
                'html', 'css', 'sass', 'scss', 'less', 'bootstrap', 'tailwind css', 'bulma',
                'react', 'angular', 'vue.js', 'svelte', 'ember.js', 'backbone.js', 'jquery',
                'node.js', 'express.js', 'next.js', 'nuxt.js', 'gatsby', 'webpack', 'vite',
                'parcel', 'rollup', 'npm', 'yarn', 'pnpm', 'rest api', 'graphql', 'apollo',
                'websockets', 'websocket', 'ajax', 'json', 'xml', 'soap'
            ],
            'backend_frameworks': [
                'django', 'flask', 'fastapi', 'tornado', 'pyramid', 'spring boot', 'spring mvc',
                'laravel', 'symfony', 'codeigniter', 'ruby on rails', 'sinatra', 'asp.net',
                '.net core', 'express.js', 'koa.js', 'nestjs', 'adonis.js', 'meteor'
            ],
            'databases': [
                'mysql', 'postgresql', 'sqlite', 'oracle', 'sql server', 'mongodb', 'cassandra',
                'redis', 'elasticsearch', 'neo4j', 'couchdb', 'dynamodb', 'firebase firestore',
                'firebase realtime database', 'influxdb', 'clickhouse', 'snowflake', 'bigquery',
                'amazon rds', 'amazon aurora', 'mariadb', 'cockroachdb', 'prisma', 'typeorm',
                'sequelize', 'mongoose', 'sql', 'nosql', 'acid', 'cap theorem'
            ],
            'cloud_platforms': [
                'aws', 'amazon web services', 'azure', 'microsoft azure', 'google cloud platform',
                'gcp', 'google cloud', 'alibaba cloud', 'ibm cloud', 'oracle cloud', 'digitalocean',
                'linode', 'vultr', 'heroku', 'netlify', 'vercel', 'cloudflare', 'firebase'
            ],
            'cloud_services': [
                'ec2', 's3', 'rds', 'lambda', 'cloudformation', 'cloudwatch', 'iam', 'vpc',
                'elastic beanstalk', 'ecs', 'eks', 'fargate', 'api gateway', 'cognito',
                'azure functions', 'azure storage', 'azure sql', 'azure cosmos db',
                'compute engine', 'cloud storage', 'cloud functions', 'cloud run',
                'kubernetes engine', 'app engine', 'cloud sql', 'pub/sub'
            ],
            'devops_tools': [
                'docker', 'kubernetes', 'jenkins', 'gitlab ci', 'github actions', 'travis ci',
                'circleci', 'ansible', 'terraform', 'chef', 'puppet', 'vagrant', 'helm',
                'prometheus', 'grafana', 'elk stack', 'elasticsearch', 'logstash', 'kibana',
                'nagios', 'zabbix', 'datadog', 'new relic', 'splunk', 'nginx', 'apache',
                'load balancing', 'reverse proxy', 'microservices', 'serverless', 'istio'
            ],
            'version_control': [
                'git', 'github', 'gitlab', 'bitbucket', 'svn', 'mercurial', 'perforce',
                'git flow', 'github flow', 'branching strategy', 'pull request', 'merge request'
            ],
            'data_science_ml': [
                'machine learning', 'deep learning', 'artificial intelligence', 'neural networks',
                'natural language processing', 'nlp', 'computer vision', 'reinforcement learning',
                'supervised learning', 'unsupervised learning', 'classification', 'regression',
                'clustering', 'dimensionality reduction', 'feature engineering', 'model selection',
                'cross validation', 'hyperparameter tuning', 'ensemble methods', 'random forest',
                'gradient boosting', 'xgboost', 'lightgbm', 'catboost', 'svm', 'support vector machine',
                'decision trees', 'naive bayes', 'k-means', 'hierarchical clustering', 'dbscan',
                'pca', 'principal component analysis', 'linear regression', 'logistic regression'
            ],
            'ml_frameworks': [
                'tensorflow', 'pytorch', 'keras', 'scikit-learn', 'pandas', 'numpy', 'scipy',
                'matplotlib', 'seaborn', 'plotly', 'bokeh', 'opencv', 'pillow', 'scikit-image',
                'nltk', 'spacy', 'gensim', 'transformers', 'hugging face', 'fastai', 'jax',
                'mlflow', 'kubeflow', 'weights and biases', 'wandb', 'tensorboard'
            ],
            'data_tools': [
                'tableau', 'power bi', 'qlikview', 'looker', 'superset', 'metabase', 'grafana',
                'jupyter', 'jupyter notebook', 'jupyter lab', 'google colab', 'kaggle',
                'databricks', 'apache spark', 'pyspark', 'hadoop', 'hive', 'pig', 'hbase',
                'kafka', 'airflow', 'luigi', 'prefect', 'dbt', 'great expectations'
            ],
            'mobile_development': [
                'android', 'ios', 'react native', 'flutter', 'xamarin', 'ionic', 'cordova',
                'phonegap', 'native script', 'unity', 'unreal engine', 'android studio',
                'xcode', 'java android', 'kotlin android', 'swift ios', 'objective-c ios'
            ],
            'testing': [
                'unit testing', 'integration testing', 'end-to-end testing', 'selenium', 'cypress',
                'jest', 'mocha', 'chai', 'jasmine', 'karma', 'protractor', 'playwright',
                'testng', 'junit', 'pytest', 'unittest', 'mockito', 'sinon', 'enzyme',
                'react testing library', 'postman', 'insomnia', 'newman', 'k6', 'locust'
            ],
            'project_management': [
                'agile', 'scrum', 'kanban', 'waterfall', 'lean', 'six sigma', 'prince2',
                'pmp', 'jira', 'confluence', 'trello', 'asana', 'monday.com', 'notion',
                'slack', 'microsoft teams', 'zoom', 'miro', 'figma', 'sketch', 'invision'
            ],
            'security': [
                'cybersecurity', 'information security', 'network security', 'web security',
                'application security', 'penetration testing', 'vulnerability assessment',
                'ethical hacking', 'ceh', 'cissp', 'cism', 'cisa', 'owasp', 'ssl', 'tls',
                'encryption', 'cryptography', 'oauth', 'jwt', 'saml', 'ldap', 'active directory'
            ],
            'design_tools': [
                'adobe photoshop', 'adobe illustrator', 'adobe xd', 'figma', 'sketch',
                'invision', 'principle', 'framer', 'zeplin', 'marvel', 'canva', 'gimp',
                'inkscape', 'blender', 'maya', 'autocad', 'solidworks', 'fusion 360'
            ],
            'soft_skills': [
                'leadership', 'communication', 'teamwork', 'collaboration', 'problem solving',
                'analytical thinking', 'critical thinking', 'creative thinking', 'time management',
                'project management', 'adaptability', 'flexibility', 'mentoring', 'coaching',
                'presentation skills', 'public speaking', 'negotiation', 'conflict resolution',
                'decision making', 'strategic thinking', 'emotional intelligence', 'empathy'
            ],
            'certifications': [
                'aws certified', 'azure certified', 'google cloud certified', 'comptia',
                'cisco certified', 'microsoft certified', 'oracle certified', 'salesforce certified',
                'pmp', 'scrum master', 'product owner', 'itil', 'lean six sigma', 'ceh',
                'cissp', 'cism', 'cisa', 'togaf', 'cobit', 'prince2'
            ]
        }
    
    def extract_text_from_file(self, uploaded_file) -> str:
        """Enhanced text extraction with better error handling"""
        try:
            file_bytes = uploaded_file.read()
            
            if uploaded_file.type == "application/pdf":
                return self._extract_text_from_pdf(file_bytes)
            elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                return self._extract_text_from_docx(file_bytes)
            else:
                raise ValueError(f"Unsupported file type: {uploaded_file.type}")
                
        except Exception as e:
            st.error(f"❌ Error extracting text from file: {str(e)}")
            return ""
    
    def _extract_text_from_pdf(self, file_bytes: bytes) -> str:
        """Enhanced PDF text extraction"""
        try:
            pdf_document = fitz.open(stream=file_bytes, filetype="pdf")
            text = ""
            
            for page_num in range(pdf_document.page_count):
                page = pdf_document[page_num]
                page_text = page.get_text()
                
                # Try different extraction methods if text is minimal
                if len(page_text.strip()) < 50:
                    # Try text blocks
                    blocks = page.get_text("blocks")
                    for block in blocks:
                        if isinstance(block, tuple) and len(block) >= 5:
                            text += block[4] + "\n"
                else:
                    text += page_text + "\n"
            
            pdf_document.close()
            return self._clean_and_normalize_text(text)
            
        except Exception as e:
            st.error(f"❌ Error reading PDF: {str(e)}")
            return ""
    
    def _extract_text_from_docx(self, file_bytes: bytes) -> str:
        """Enhanced DOCX text extraction"""
        try:
            doc = docx.Document(BytesIO(file_bytes))
            text = ""
            
            # Extract from paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return self._clean_and_normalize_text(text)
            
        except Exception as e:
            st.error(f"❌ Error reading DOCX: {str(e)}")
            return ""
    
    def _clean_and_normalize_text(self, text: str) -> str:
        """Enhanced text cleaning and normalization"""
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'\n+', '\n', text)
        
        # Remove special characters but keep important ones
        text = re.sub(r'[^\w\s\.\,\-\@\+\(\)\[\]\:\;\/\\]', ' ', text)
        
        # Normalize common abbreviations
        text = re.sub(r'\bB\.?Tech\b', 'BTech', text, flags=re.IGNORECASE)
        text = re.sub(r'\bM\.?Tech\b', 'MTech', text, flags=re.IGNORECASE)
        text = re.sub(r'\bB\.?Sc\b', 'BSc', text, flags=re.IGNORECASE)
        text = re.sub(r'\bM\.?Sc\b', 'MSc', text, flags=re.IGNORECASE)
        
        # Fix common OCR errors
        text = re.sub(r'\bl\b', 'I', text)  # lowercase l to uppercase I
        text = re.sub(r'\b0\b', 'O', text)  # zero to letter O in names
        
        return text.strip()
    
    def parse_resume(self, resume_text: str) -> Dict:
        """Enhanced resume parsing with better structure"""
        if not resume_text.strip():
            return self._get_empty_resume_structure()
        
        parsed_data = {
            'contact_info': self._extract_contact_info(resume_text),
            'personal_info': self._extract_personal_info(resume_text),
            'skills': self._extract_skills_comprehensive(resume_text),
            'experience': self._extract_experience_detailed(resume_text),
            'education': self._extract_education_detailed(resume_text),
            'projects': self._extract_projects_detailed(resume_text),
            'certifications': self._extract_certifications_detailed(resume_text),
            'achievements': self._extract_achievements(resume_text),
            'languages': self._extract_languages(resume_text),
            'summary': self._extract_professional_summary(resume_text)
        }
        
        return parsed_data
    
    def _get_empty_resume_structure(self) -> Dict:
        """Return empty resume structure"""
        return {
            'contact_info': {},
            'personal_info': {},
            'skills': {},
            'experience': [],
            'education': [],
            'projects': [],
            'certifications': [],
            'achievements': [],
            'languages': [],
            'summary': ''
        }
    
    def _extract_contact_info(self, text: str) -> Dict:
        """Enhanced contact information extraction"""
        contact_info = {}
        
        # Email - multiple patterns
        email_patterns = [
            r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b',
            r'email\s*:?\s*([A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,})',
            r'e-mail\s*:?\s*([A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,})'
        ]
        
        for pattern in email_patterns:
            emails = re.findall(pattern, text, re.IGNORECASE)
            if emails:
                contact_info['email'] = emails[0] if isinstance(emails[0], str) else emails[0]
                break
        
        # Phone - multiple patterns for Indian and international numbers
        phone_patterns = [
            r'(?:\+91[-.\s]?)?(?:\(?\d{3}\)?[-.\s]?)?\d{3}[-.\s]?\d{4}',
            r'(?:\+91\s?)?[6-9]\d{9}',
            r'phone\s*:?\s*((?:\+91[-.\s]?)?(?:\(?\d{3}\)?[-.\s]?)?\d{3}[-.\s]?\d{4})',
            r'mobile\s*:?\s*((?:\+91[-.\s]?)?(?:\(?\d{3}\)?[-.\s]?)?\d{3}[-.\s]?\d{4})',
            r'contact\s*:?\s*((?:\+91[-.\s]?)?(?:\(?\d{3}\)?[-.\s]?)?\d{3}[-.\s]?\d{4})'
        ]
        
        for pattern in phone_patterns:
            phones = re.findall(pattern, text, re.IGNORECASE)
            if phones:
                contact_info['phone'] = phones[0] if isinstance(phones[0], str) else phones[0]
                break
        
        # LinkedIn
        linkedin_patterns = [
            r'linkedin\.com/in/[\w-]+',
            r'linkedin\s*:?\s*([\w-]+)',
            r'linkedin\.com/in/([\w-]+)'
        ]
        
        for pattern in linkedin_patterns:
            linkedin = re.findall(pattern, text, re.IGNORECASE)
            if linkedin:
                profile = linkedin[0]
                if 'linkedin.com' not in profile:
                    profile = f"linkedin.com/in/{profile}"
                contact_info['linkedin'] = profile
                break
        
        # GitHub
        github_patterns = [
            r'github\.com/[\w-]+',
            r'github\s*:?\s*([\w-]+)',
            r'github\.com/([\w-]+)'
        ]
        
        for pattern in github_patterns:
            github = re.findall(pattern, text, re.IGNORECASE)
            if github:
                profile = github[0]
                if 'github.com' not in profile:
                    profile = f"github.com/{profile}"
                contact_info['github'] = profile
                break
        
        # Location/Address
        location_patterns = [
            r'(?:address|location|city)\s*:?\s*([^,\n]+(?:,\s*[^,\n]+)*)',
            r'([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*,\s*[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)',
            r'(\b(?:Hyderabad|Bangalore|Mumbai|Delhi|Chennai|Pune|Kolkata|Ahmedabad|Jaipur|Lucknow)\b[^,\n]*)'
        ]
        
        for pattern in location_patterns:
            locations = re.findall(pattern, text, re.IGNORECASE)
            if locations:
                contact_info['location'] = locations[0].strip()
                break
        
        return contact_info
    
    def _extract_personal_info(self, text: str) -> Dict:
        """Extract personal information"""
        personal_info = {}
        
        # Extract name (usually at the beginning)
        lines = text.split('\n')
        for line in lines[:5]:  # Check first 5 lines
            line = line.strip()
            if len(line) > 3 and len(line) < 50:
                # Check if it looks like a name
                if re.match(r'^[A-Z][a-z]+(?:\s+[A-Z][a-z]+)+$', line):
                    personal_info['name'] = line
                    break
        
        return personal_info
    
    def _extract_skills_comprehensive(self, text: str) -> Dict[str, List[str]]:
        """Enhanced skills extraction with categorization"""
        text_lower = text.lower()
        extracted_skills = {}
        
        # Enhanced matching with context
        for category, skills_list in self.skills_keywords.items():
            found_skills = []
            for skill in skills_list:
                # Create flexible pattern
                skill_pattern = re.escape(skill.lower()).replace(r'\ ', r'[\s\-_]*')
                pattern = r'\b' + skill_pattern + r'\b'
                
                if re.search(pattern, text_lower):
                    found_skills.append(skill)
            
            if found_skills:
                extracted_skills[category] = sorted(list(set(found_skills)))
        
        # Extract skills from skills section specifically
        skills_section = self._extract_skills_section(text)
        if skills_section:
            additional_skills = self._parse_skills_section(skills_section)
            for category, skills in additional_skills.items():
                if category in extracted_skills:
                    extracted_skills[category].extend(skills)
                    extracted_skills[category] = sorted(list(set(extracted_skills[category])))
                else:
                    extracted_skills[category] = skills
        
        return extracted_skills
    
    def _extract_skills_section(self, text: str) -> str:
        """Extract dedicated skills section"""
        patterns = [
            r'(?:technical\s+)?skills?\s*:?\s*(.*?)(?=\n\s*(?:experience|education|projects|certifications|achievements|summary|objective)|$)',
            r'(?:core\s+)?competencies\s*:?\s*(.*?)(?=\n\s*(?:experience|education|projects|certifications|achievements|summary|objective)|$)',
            r'technologies\s*:?\s*(.*?)(?=\n\s*(?:experience|education|projects|certifications|achievements|summary|objective)|$)',
            r'expertise\s*:?\s*(.*?)(?=\n\s*(?:experience|education|projects|certifications|achievements|summary|objective)|$)'
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
            if match:
                return match.group(1).strip()
        
        return ""
    
    def _parse_skills_section(self, skills_text: str) -> Dict[str, List[str]]:
        """Parse skills from dedicated skills section"""
        skills_dict = {}
        
        # Split by common separators
        skills_items = re.split(r'[,;\n\|•\*\-]', skills_text)
        
        for item in skills_items:
            item = item.strip()
            if len(item) > 2:
                # Categorize based on known skills
                categorized = False
                for category, known_skills in self.skills_keywords.items():
                    for known_skill in known_skills:
                        if known_skill.lower() in item.lower():
                            if category not in skills_dict:
                                skills_dict[category] = []
                            skills_dict[category].append(item)
                            categorized = True
                            break
                    if categorized:
                        break
                
                # If not categorized, add to general skills
                if not categorized and len(item) < 30:
                    if 'other_skills' not in skills_dict:
                        skills_dict['other_skills'] = []
                    skills_dict['other_skills'].append(item)
        
        return skills_dict
    
    def _extract_experience_detailed(self, text: str) -> List[Dict]:
        """Enhanced experience extraction"""
        experience = []
        
        # Look for experience section
        experience_patterns = [
            r'(?:professional\s+)?(?:work\s+)?experience\s*:?\s*(.*?)(?=\n\s*(?:education|skills|projects|certifications|achievements|summary|objective)|$)',
            r'(?:employment\s+)?history\s*:?\s*(.*?)(?=\n\s*(?:education|skills|projects|certifications|achievements|summary|objective)|$)',
            r'career\s+summary\s*:?\s*(.*?)(?=\n\s*(?:education|skills|projects|certifications|achievements|summary|objective)|$)'
        ]
        
        experience_text = ""
        for pattern in experience_patterns:
            match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
            if match:
                experience_text = match.group(1).strip()
                break
        
        if experience_text:
            # Parse individual jobs
            job_entries = self._parse_job_entries(experience_text)
            experience.extend(job_entries)
        
        # Also extract years mentioned throughout resume
        years_experience = self._extract_total_experience(text)
        if years_experience:
            experience.append({
                'type': 'total_experience',
                'years': years_experience,
                'source': 'mentioned_in_resume'
            })
        
        return experience
    
    def _parse_job_entries(self, experience_text: str) -> List[Dict]:
        """Parse individual job entries"""
        jobs = []
        
        # Split by potential job boundaries
        job_sections = re.split(r'\n(?=\w+.*?(?:analyst|engineer|developer|manager|specialist|consultant|intern|associate))', experience_text, flags=re.IGNORECASE)
        
        for section in job_sections:
            if len(section.strip()) < 20:
                continue
                
            job_info = {}
            
            # Extract job title
            title_match = re.search(r'^([^\n]+)', section.strip())
            if title_match:
                job_info['title'] = title_match.group(1).strip()
            
            # Extract company
            company_patterns = [
                r'@\s*([^\n,]+)',
                r'at\s+([^\n,]+)',
                r'company\s*:?\s*([^\n,]+)'
            ]
            
            for pattern in company_patterns:
                company_match = re.search(pattern, section, re.IGNORECASE)
                if company_match:
                    job_info['company'] = company_match.group(1).strip()
                    break
            
            # Extract duration
            duration_patterns = [
                r'(\d{4})\s*[-–]\s*(\d{4}|present)',
                r'(\w+\s+\d{4})\s*[-–]\s*(\w+\s+\d{4}|present)',
                r'(\d+)\s*years?\s*(\d+)?\s*months?'
            ]
            
            for pattern in duration_patterns:
                duration_match = re.search(pattern, section, re.IGNORECASE)
                if duration_match:
                    job_info['duration'] = duration_match.group(0)
                    break
            
            # Extract key responsibilities/achievements
            lines = section.split('\n')[1:]  # Skip title line
            responsibilities = []
            for line in lines:
                line = line.strip()
                if line and len(line) > 20:
                    # Remove bullet points
                    line = re.sub(r'^[•\*\-\+]\s*', '', line)
                    responsibilities.append(line)
            
            job_info['responsibilities'] = responsibilities[:5]  # Top 5
            
            if job_info:
                jobs.append(job_info)
        
        return jobs
    
    def _extract_total_experience(self, text: str) -> int:
        """Extract total years of experience mentioned"""
        patterns = [
            r'(\d+)\+?\s*years?\s+(?:of\s+)?(?:total\s+)?(?:professional\s+)?experience',
            r'(?:total\s+)?(?:professional\s+)?experience\s*:?\s*(\d+)\+?\s*years?',
            r'(\d+)\+?\s*years?\s+(?:in\s+)?(?:software|development|engineering|data|marketing|sales)',
            r'over\s+(\d+)\s*years?\s+(?:of\s+)?experience'
        ]
        
        max_years = 0
        for pattern in patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            for match in matches:
                try:
                    years = int(match)
                    max_years = max(max_years, years)
                except ValueError:
                    continue
        
        return max_years
    
    def _extract_education_detailed(self, text: str) -> List[Dict]:
        """Enhanced education extraction"""
        education = []
        
        # Education patterns
        education_patterns = [
            r'education\s*:?\s*(.*?)(?=\n\s*(?:experience|skills|projects|certifications|achievements|summary|objective)|$)',
            r'academic\s+(?:background|qualifications)\s*:?\s*(.*?)(?=\n\s*(?:experience|skills|projects|certifications|achievements|summary|objective)|$)',
            r'qualifications\s*:?\s*(.*?)(?=\n\s*(?:experience|skills|projects|certifications|achievements|summary|objective)|$)'
        ]
        
        education_text = ""
        for pattern in education_patterns:
            match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
            if match:
                education_text = match.group(1).strip()
                break
        
        # Degree patterns
        degree_patterns = [
            r'(bachelor\'?s?|b\.?tech|b\.?sc|b\.?com|b\.?ba|bba|bca|be|b\.?e\.?)\s*(?:of|in)?\s*([^\n,]+?)(?:\s*(?:from|at)\s*([^\n,]+?))?(?:\s*[-–]\s*(\d{4}))?',
            r'(master\'?s?|m\.?tech|m\.?sc|m\.?com|m\.?ba|mba|mca|me|m\.?e\.?)\s*(?:of|in)?\s*([^\n,]+?)(?:\s*(?:from|at)\s*([^\n,]+?))?(?:\s*[-–]\s*(\d{4}))?',
            r'(phd|doctorate|ph\.?d\.?)\s*(?:of|in)?\s*([^\n,]+?)(?:\s*(?:from|at)\s*([^\n,]+?))?(?:\s*[-–]\s*(\d{4}))?',
            r'(diploma|certificate)\s*(?:of|in)?\s*([^\n,]+?)(?:\s*(?:from|at)\s*([^\n,]+?))?(?:\s*[-–]\s*(\d{4}))?'
        ]
        
        # Search in education section first
        if education_text:
            for pattern in degree_patterns:
                matches = re.findall(pattern, education_text, re.IGNORECASE)
                for match in matches:
                    education.append({
                        'degree': match[0].strip(),
                        'field': match[1].strip() if len(match) > 1 and match[1] else '',
                        'institution': match[2].strip() if len(match) > 2 and match[2] else '',
                        'year': match[3] if len(match) > 3 and match[3] else ''
                    })
        
        # Search entire document if no education section found
        if not education:
            for pattern in degree_patterns:
                matches = re.findall(pattern, text, re.IGNORECASE)
                for match in matches:
                    education.append({
                        'degree': match[0].strip(),
                        'field': match[1].strip() if len(match) > 1 and match[1] else '',
                        'institution': match[2].strip() if len(match) > 2 and match[2] else '',
                        'year': match[3] if len(match) > 3 and match[3] else ''
                    })
        
        return education[:5]  # Top 5 education entries
    
    def _extract_projects_detailed(self, text: str) -> List[Dict]:
        """Enhanced project extraction"""
        projects = []
        
        # Project section patterns
        project_patterns = [
            r'projects?\s*:?\s*(.*?)(?=\n\s*(?:experience|education|skills|certifications|achievements|summary|objective)|$)',
            r'(?:key\s+)?(?:academic\s+)?projects?\s*:?\s*(.*?)(?=\n\s*(?:experience|education|skills|certifications|achievements|summary|objective)|$)',
            r'portfolio\s*:?\s*(.*?)(?=\n\s*(?:experience|education|skills|certifications|achievements|summary|objective)|$)'
        ]
        
        project_text = ""
        for pattern in project_patterns:
            match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
            if match:
                project_text = match.group(1).strip()
                break
        
        if project_text:
            # Parse individual projects
            project_entries = re.split(r'\n(?=\d+\.|[•\*\-]\s*\w+|\w+\s*[-–:])', project_text)
            
            for entry in project_entries:
                if len(entry.strip()) < 20:
                    continue
                
                project_info = {}
                lines = entry.strip().split('\n')
                
                # First line is usually the title
                title_line = lines[0].strip()
                title_line = re.sub(r'^\d+\.\s*|^[•\*\-]\s*', '', title_line)
                project_info['title'] = title_line[:100]
                
                # Extract description from remaining lines
                description_lines = []
                for line in lines[1:]:
                    line = line.strip()
                    if line:
                        line = re.sub(r'^[•\*\-]\s*', '', line)
                        description_lines.append(line)
                
                project_info['description'] = ' '.join(description_lines)[:300]
                
                # Extract technologies used
                tech_pattern = r'(?:technologies?|tools?|stack|built\s+(?:with|using))\s*:?\s*([^\n]+)'
                tech_match = re.search(tech_pattern, entry, re.IGNORECASE)
                if tech_match:
                    technologies = [tech.strip() for tech in re.split(r'[,;\|]', tech_match.group(1))]
                    project_info['technologies'] = technologies[:10]
                
                # Extract links
                link_patterns = [
                    r'(?:github|git)\s*:?\s*(https?://[^\s]+)',
                    r'(?:demo|live|url|link)\s*:?\s*(https?://[^\s]+)',
                    r'(https?://github\.com/[^\s]+)',
                    r'(https?://[^\s]+\.(?:com|org|net|io)/[^\s]*)'
                ]
                
                for pattern in link_patterns:
                    link_match = re.search(pattern, entry, re.IGNORECASE)
                    if link_match:
                        project_info['link'] = link_match.group(1)
                        break
                
                if project_info:
                    projects.append(project_info)
        
        return projects[:10]  # Top 10 projects
    
    def _extract_certifications_detailed(self, text: str) -> List[Dict]:
        """Enhanced certification extraction"""
        certifications = []
        
        # Certification section patterns
        cert_patterns = [
            r'certifications?\s*:?\s*(.*?)(?=\n\s*(?:experience|education|skills|projects|achievements|summary|objective)|$)',
            r'(?:licenses?\s+and\s+)?certifications?\s*:?\s*(.*?)(?=\n\s*(?:experience|education|skills|projects|achievements|summary|objective)|$)',
            r'professional\s+certifications?\s*:?\s*(.*?)(?=\n\s*(?:experience|education|skills|projects|achievements|summary|objective)|$)'
        ]
        
        cert_text = ""
        for pattern in cert_patterns:
            match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
            if match:
                cert_text = match.group(1).strip()
                break
        
        # Known certification providers and their certifications
        cert_keywords = {
            'aws': ['aws certified', 'amazon web services', 'cloud practitioner', 'solutions architect', 'developer associate'],
            'azure': ['azure certified', 'microsoft azure', 'azure fundamentals', 'azure administrator'],
            'google cloud': ['google cloud certified', 'gcp certified', 'cloud engineer', 'data engineer'],
            'cisco': ['cisco certified', 'ccna', 'ccnp', 'ccie'],
            'microsoft': ['microsoft certified', 'mcsa', 'mcse', 'office specialist'],
            'oracle': ['oracle certified', 'ocp', 'oca'],
            'comptia': ['comptia', 'security+', 'network+', 'a+'],
            'project management': ['pmp', 'prince2', 'scrum master', 'product owner'],
            'security': ['cissp', 'ceh', 'cism', 'cisa'],
            'data science': ['certified data scientist', 'machine learning certification', 'analytics certification']
        }
        
        # Search in certification section
        if cert_text:
            cert_entries = re.split(r'\n|[•\*\-]', cert_text)
            for entry in cert_entries:
                entry = entry.strip()
                if len(entry) > 5:
                    cert_info = {'name': entry}
                    
                    # Try to identify provider
                    for provider, keywords in cert_keywords.items():
                        if any(keyword.lower() in entry.lower() for keyword in keywords):
                            cert_info['provider'] = provider
                            break
                    
                    # Extract year
                    year_match = re.search(r'(\d{4})', entry)
                    if year_match:
                        cert_info['year'] = year_match.group(1)
                    
                    certifications.append(cert_info)
        
        # Also search entire document for common certifications
        for provider, keywords in cert_keywords.items():
            for keyword in keywords:
                pattern = r'[^\n]*' + re.escape(keyword.lower()) + r'[^\n]*'
                matches = re.findall(pattern, text.lower())
                
                for match in matches:
                    if len(match.strip()) > 10 and match.strip() not in [c['name'] for c in certifications]:
                        cert_info = {
                            'name': match.strip()[:100],
                            'provider': provider
                        }
                        
                        # Extract year
                        year_match = re.search(r'(\d{4})', match)
                        if year_match:
                            cert_info['year'] = year_match.group(1)
                        
                        certifications.append(cert_info)
        
        return certifications[:10]  # Top 10 certifications
    
    def _extract_achievements(self, text: str) -> List[str]:
        """Extract achievements and awards"""
        achievements = []
        
        # Achievement section patterns
        achievement_patterns = [
            r'achievements?\s*:?\s*(.*?)(?=\n\s*(?:experience|education|skills|projects|certifications|summary|objective)|$)',
            r'awards?\s*:?\s*(.*?)(?=\n\s*(?:experience|education|skills|projects|certifications|summary|objective)|$)',
            r'honors?\s*:?\s*(.*?)(?=\n\s*(?:experience|education|skills|projects|certifications|summary|objective)|$)',
            r'accomplishments?\s*:?\s*(.*?)(?=\n\s*(?:experience|education|skills|projects|certifications|summary|objective)|$)'
        ]
        
        for pattern in achievement_patterns:
            match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
            if match:
                achievement_text = match.group(1).strip()
                achievement_entries = re.split(r'\n|[•\*\-]', achievement_text)
                
                for entry in achievement_entries:
                    entry = entry.strip()
                    if len(entry) > 10:
                        achievements.append(entry[:150])
        
        return achievements[:8]  # Top 8 achievements
    
    def _extract_languages(self, text: str) -> List[Dict]:
        """Extract language skills"""
        languages = []
        
        # Language patterns
        language_patterns = [
            r'languages?\s*:?\s*(.*?)(?=\n\s*(?:experience|education|skills|projects|certifications|achievements|summary|objective)|$)',
            r'linguistic\s+skills?\s*:?\s*(.*?)(?=\n\s*(?:experience|education|skills|projects|certifications|achievements|summary|objective)|$)'
        ]
        
        common_languages = [
            'english', 'hindi', 'spanish', 'french', 'german', 'chinese', 'japanese', 'korean',
            'arabic', 'russian', 'portuguese', 'italian', 'dutch', 'swedish', 'telugu', 'tamil',
            'bengali', 'marathi', 'gujarati', 'punjabi', 'kannada', 'malayalam', 'urdu'
        ]
        
        proficiency_levels = ['native', 'fluent', 'advanced', 'intermediate', 'basic', 'beginner']
        
        for pattern in language_patterns:
            match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
            if match:
                language_text = match.group(1).strip()
                
                for language in common_languages:
                    if language.lower() in language_text.lower():
                        lang_info = {'language': language.title()}
                        
                        # Try to find proficiency level
                        for level in proficiency_levels:
                            if level.lower() in language_text.lower():
                                lang_info['proficiency'] = level.title()
                                break
                        
                        languages.append(lang_info)
        
        return languages[:8]  # Top 8 languages
    
    def _extract_professional_summary(self, text: str) -> str:
        """Extract professional summary or objective"""
        summary_patterns = [
            r'(?:professional\s+)?summary\s*:?\s*(.*?)(?=\n\s*(?:experience|education|skills|projects|certifications|achievements)|$)',
            r'objective\s*:?\s*(.*?)(?=\n\s*(?:experience|education|skills|projects|certifications|achievements)|$)',
            r'profile\s*:?\s*(.*?)(?=\n\s*(?:experience|education|skills|projects|certifications|achievements)|$)',
            r'about\s+me\s*:?\s*(.*?)(?=\n\s*(?:experience|education|skills|projects|certifications|achievements)|$)'
        ]
        
        for pattern in summary_patterns:
            match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
            if match:
                summary = match.group(1).strip()
                # Clean up the summary
                summary = re.sub(r'\s+', ' ', summary)
                return summary[:500]  # Limit to 500 characters
        
        return ""
    
    def calculate_experience_years(self, parsed_resume: Dict) -> int:
        """Calculate total years of experience from parsed resume"""
        max_years = 0
        
        # Check direct experience mention
        experience_data = parsed_resume.get('experience', [])
        for exp in experience_data:
            if exp.get('type') == 'total_experience':
                max_years = max(max_years, exp.get('years', 0))
            
            # Parse duration from job entries
            duration = exp.get('duration', '')
            if duration:
                # Extract years from duration
                year_matches = re.findall(r'(\d+)\s*years?', duration.lower())
                for year_match in year_matches:
                    max_years = max(max_years, int(year_match))
                
                # Parse year ranges
                range_matches = re.findall(r'(\d{4})\s*[-–]\s*(\d{4}|present)', duration.lower())
                for start, end in range_matches:
                    end_year = 2024 if end == 'present' else int(end)
                    years = end_year - int(start)
                    max_years = max(max_years, years)
        
        # If no experience found, estimate from skills count
        if max_years == 0:
            total_skills = sum(len(skills) for skills in parsed_resume.get('skills', {}).values())
            if total_skills > 30:
                max_years = 3
            elif total_skills > 15:
                max_years = 2
            elif total_skills > 8:
                max_years = 1
        
        return max_years
    
    def get_resume_strength_score(self, parsed_resume: Dict) -> int:
        """Calculate overall resume strength score"""
        score = 0
        
        # Contact information (20 points)
        contact_info = parsed_resume.get('contact_info', {})
        if contact_info.get('email'):
            score += 10
        if contact_info.get('phone'):
            score += 5
        if contact_info.get('linkedin'):
            score += 3
        if contact_info.get('github'):
            score += 2
        
        # Skills (25 points)
        skills = parsed_resume.get('skills', {})
        skill_count = sum(len(skill_list) for skill_list in skills.values())
        score += min(25, skill_count * 2)
        
        # Experience (25 points)
        experience = parsed_resume.get('experience', [])
        if experience:
            score += min(25, len(experience) * 5)
        
        # Education (15 points)
        education = parsed_resume.get('education', [])
        if education:
            score += min(15, len(education) * 8)
        
        # Projects (10 points)
        projects = parsed_resume.get('projects', [])
        if projects:
            score += min(10, len(projects) * 2)
        
        # Certifications (5 points)
        certifications = parsed_resume.get('certifications', [])
        if certifications:
            score += min(5, len(certifications))
        
        return min(100, score)
                